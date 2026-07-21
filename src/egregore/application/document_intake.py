from __future__ import annotations

import hashlib
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from egregore.interface.ports.dossier_ports import DossierGenerateRequest
from egregore.shared.canonical import canonical_dumps


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    content_type: str
    text: str
    fingerprint: str


def _sha256_hex(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError("pdfplumber is not installed") from exc

    parts: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    try:
        import docx
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is not installed") from exc

    document = docx.Document(BytesIO(file_bytes))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    return "\n".join(parts)


def extract_text_from_odt(file_bytes: bytes) -> str:
    """Extract plain text from ODT bytes using odfpy."""
    try:
        from odf import opendocument
        from odf.text import P
    except ModuleNotFoundError as exc:  # pragma: no cover
        raise RuntimeError("odfpy is not installed") from exc

    doc = opendocument.load(BytesIO(file_bytes))

    def _get_text(node) -> str:
        parts: list[str] = []
        for child in node.childNodes:
            if hasattr(child, "data"):
                parts.append(child.data)
            elif hasattr(child, "childNodes"):
                parts.append(_get_text(child))
        return "".join(parts)

    paragraphs = doc.getElementsByType(P)
    texts: list[str] = []
    for p in paragraphs:
        txt = _get_text(p).strip()
        if txt:
            texts.append(txt)
    return "\n".join(texts)


def _classify_document(filename: str, text: str) -> dict[str, Any]:  # noqa: C901
    """
    Lightweight deterministic classification based on filename and content heuristics.
    No ML — pure rules suitable for compliance demos.
    """
    fname_lower = filename.lower()
    text_lower = text.lower()

    doc_type = "unknown"
    jurisdiction = "ca-federal"
    risk_flags: list[str] = []

    # Filename heuristics
    if any(k in fname_lower for k in ("sop", "standard", "procedure")):
        doc_type = "sop"
    elif any(k in fname_lower for k in ("license", "permit", "authorization")):
        doc_type = "license"
    elif any(k in fname_lower for k in ("batch", "lot", "test", "lab")):
        doc_type = "lab_report"
    elif any(k in fname_lower for k in ("seed", "sale", "manifest", "shipping")):
        doc_type = "seed_to_sale_manifest"
    elif any(k in fname_lower for k in ("employee", "screening", "background")):
        doc_type = "employee_screening"

    # Content heuristics (override filename if strong signal)
    if "seed-to-sale" in text_lower or "manifest" in text_lower:
        doc_type = "seed_to_sale_manifest"
    elif "standard operating procedure" in text_lower:
        doc_type = "sop"
    elif "health canada" in text_lower or "santé canada" in text_lower:
        jurisdiction = "ca-health-canada"
    elif "tax" in text_lower or "revenue" in text_lower or "cra" in text_lower:
        doc_type = "tax_reconciliation"
        jurisdiction = "ca-cra"

    # Risk flags
    if "expired" in text_lower or "lapsed" in text_lower:
        risk_flags.append("expired_document")
    if "non-compliant" in text_lower or "violation" in text_lower:
        risk_flags.append("compliance_violation")
    if len(text.strip()) < 50:
        risk_flags.append("low_content")

    return {
        "doc_type": doc_type,
        "jurisdiction": jurisdiction,
        "risk_flags": risk_flags,
        "filename": filename,
        "text_preview": text[:500],
    }


def extract_document(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """Route to correct extractor based on file extension."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif suffix == ".docx":
        text = extract_text_from_docx(file_bytes)
    elif suffix == ".odt":
        text = extract_text_from_odt(file_bytes)
    else:
        # Fallback: treat as plain text if readable, else base64 hint
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = f"[binary file: {filename}]"

    fingerprint = _sha256_hex(file_bytes)
    return ExtractedDocument(
        filename=filename,
        content_type=_content_type_from_suffix(suffix),
        text=text,
        fingerprint=fingerprint,
    )


def _content_type_from_suffix(suffix: str) -> str:
    mapping = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".odt": "application/vnd.oasis.opendocument.text",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    return mapping.get(suffix, "application/octet-stream")


def build_dossier_request_from_intake(
    *,
    organization_id: str,
    case_id: str,
    actor_id: str,
    causality_id: str,
    vertical: str | None,
    documents: list[ExtractedDocument],
) -> DossierGenerateRequest:
    """
    Aggregate extracted documents into a single deterministic DossierGenerateRequest.
    """
    # Deterministic canonical payload
    doc_payloads: list[dict[str, Any]] = []
    for doc in documents:
        classification = _classify_document(doc.filename, doc.text)
        doc_payloads.append(
            {
                "filename": doc.filename,
                "content_type": doc.content_type,
                "fingerprint": doc.fingerprint,
                "classification": classification,
                "extracted_text": doc.text,
            }
        )

    # Sort by filename for determinism
    doc_payloads.sort(key=lambda d: d["filename"])

    input_payload: dict[str, Any] = {
        "intake_type": "document_upload",
        "vertical": vertical,
        "documents": doc_payloads,
    }

    # Deterministic fingerprint over canonical JSON of payload
    fingerprint = _sha256_hex(canonical_dumps(input_payload).encode("utf-8"))

    return DossierGenerateRequest(
        organization_id=organization_id,
        case_id=case_id,
        actor_id=actor_id,
        input_fingerprint=fingerprint,
        engine_version="intake_v1",
        policy_version="cannabis_policy_v1",
        input_payload=input_payload,
        causality_id=causality_id,
        request_id=None,
        timestamp_ns=None,
        vertical=vertical,
    )
